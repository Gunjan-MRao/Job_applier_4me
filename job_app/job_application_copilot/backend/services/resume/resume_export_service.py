import os
from datetime import datetime, timezone

from docx import Document

from backend.services.resume.version_store import get_resume_version_by_id


def ensure_output_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


def replace_extension(filename: str, new_ext: str) -> str:
    base, _ = os.path.splitext(filename)
    return f"{base}.{new_ext}"


def export_txt(version: dict, output_dir: str) -> tuple[str, str]:
    filename = replace_extension(version["suggested_filename"], "txt")
    file_path = os.path.join(output_dir, filename)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(version["export_text"])
    return filename, file_path


def export_docx(version: dict, output_dir: str) -> tuple[str, str]:
    filename = replace_extension(version["suggested_filename"], "docx")
    file_path = os.path.join(output_dir, filename)
    doc = Document()
    for line in version["export_text"].splitlines():
        text = line.strip()
        if not text:
            doc.add_paragraph("")
            continue
        if text in {"PROFESSIONAL SUMMARY", "CORE SKILLS", "PROFESSIONAL EXPERIENCE", "EDUCATION", "ATS NOTES"}:
            doc.add_heading(text, level=1)
        elif text.startswith("- "):
            doc.add_paragraph(text[2:], style="List Bullet")
        else:
            doc.add_paragraph(text)
    doc.save(file_path)
    return filename, file_path


def export_pdf(version: dict, output_dir: str) -> tuple[str, str]:
    """PDF export — uses reportlab if available, falls back to plain text file."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas as rl_canvas

        filename = replace_extension(version["suggested_filename"], "pdf")
        file_path = os.path.join(output_dir, filename)
        c = rl_canvas.Canvas(file_path, pagesize=A4)
        width, height = A4
        x, y, lh = 50, height - 50, 16
        for raw_line in version["export_text"].splitlines():
            line = raw_line.rstrip()
            if y < 60:
                c.showPage()
                y = height - 50
            if not line:
                y -= lh
                continue
            if line in {"PROFESSIONAL SUMMARY", "CORE SKILLS", "PROFESSIONAL EXPERIENCE", "EDUCATION", "ATS NOTES"}:
                c.setFont("Helvetica-Bold", 12)
                c.drawString(x, y, line)
                y -= lh
                c.setFont("Helvetica", 10)
            else:
                c.setFont("Helvetica", 10)
                c.drawString(x, y, line[:110])
                y -= lh
        c.save()
        return filename, file_path
    except ImportError:
        # reportlab not installed — save as txt instead
        return export_txt(version, output_dir)


def export_resume_version(version_id: str, export_format: str, output_dir: str) -> dict:
    version = get_resume_version_by_id(version_id)
    if not version:
        raise ValueError("Resume version not found")
    export_format = (export_format or "").strip().lower()
    output_dir = ensure_output_dir(output_dir)
    if export_format == "txt":
        filename, file_path = export_txt(version, output_dir)
    elif export_format == "docx":
        filename, file_path = export_docx(version, output_dir)
    elif export_format == "pdf":
        filename, file_path = export_pdf(version, output_dir)
    else:
        raise ValueError("Unsupported export format. Use txt, docx, or pdf.")
    return {
        "version_id": version_id,
        "format": export_format,
        "file_path": file_path,
        "filename": filename,
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "next_action": "attach_or_review_resume_file",
    }
