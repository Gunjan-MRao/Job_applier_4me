"""backend/pipeline/exporters.py

Turn a drafted document (cover letter / cold email) into downloadable bytes.

Two pure functions with no Streamlit / network dependency so they are trivially
unit-testable and reusable by both the FastAPI backend and the Streamlit UI:

  * to_docx_bytes(text, title=...) -> bytes   (python-docx)
  * to_pdf_bytes(text, title=...)  -> bytes   (reportlab)

Both are already declared in requirements.txt (python-docx, reportlab), so no
new dependency is introduced.
"""
from __future__ import annotations

import io
import re


def _clean_lines(text: str) -> list[str]:
    return (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")


def to_docx_bytes(text: str, title: str = "") -> bytes:
    """Render ``text`` as a .docx document and return the raw file bytes."""
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for line in _clean_lines(text):
        # A blank source line becomes an empty paragraph (preserves spacing).
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf_bytes(text: str, title: str = "") -> bytes:
    """Render ``text`` as a simple, wrapped A4 PDF and return the raw bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    styles = getSampleStyleSheet()
    story = []
    if title:
        story.append(Paragraph(_escape(title), styles["Title"]))
        story.append(Spacer(1, 6 * mm))
    for line in _clean_lines(text):
        if line.strip():
            story.append(Paragraph(_escape(line), styles["BodyText"]))
        else:
            story.append(Spacer(1, 4 * mm))
    if not story:
        story.append(Paragraph("", styles["BodyText"]))
    doc.build(story)
    return buf.getvalue()


def _escape(s: str) -> str:
    """Escape reportlab mini-HTML markup so raw text renders literally."""
    return (
        s.replace("&", "&amp;")
         .replace("<", "&lt;")
         .replace(">", "&gt;")
    )


def safe_filename(*parts: str) -> str:
    """Build a filesystem-safe base filename from arbitrary parts."""
    raw = "_".join(p for p in parts if p)
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", raw).strip("_")
    return cleaned or "document"
